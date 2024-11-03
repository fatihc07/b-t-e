from flask import Flask, render_template, request, redirect, make_response, flash, jsonify
import sqlite3
from datetime import datetime
import csv
from io import BytesIO
import locale
from docx import Document
from datetime import datetime, timedelta


app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Flash mesajları için gizli anahtar

# Türkçe tarih formatı için locale ayarı
locale.setlocale(locale.LC_TIME, 'tr_TR.UTF-8')
# Veritabanı bağlantısı ve tabloların oluşturulması
def connect_db():
    conn = sqlite3.connect('budget.db')
    conn.row_factory = sqlite3.Row
    return conn

# Veritabanını başlatma ve tabloları oluşturma
def init_db():
    with connect_db() as conn:
        conn.execute('''
        CREATE TABLE IF NOT EXISTS transactions (
            id INTEGER PRIMARY KEY,
            description TEXT,
            amount REAL,
            date TEXT,
            type_id INTEGER,
            FOREIGN KEY(type_id) REFERENCES expense_types(id)
        )''')

        conn.execute('''
        CREATE TABLE IF NOT EXISTS expense_types (
            id INTEGER PRIMARY KEY,
            type_name TEXT UNIQUE
        )''')

        conn.execute('''
        CREATE TABLE IF NOT EXISTS balance (
            id INTEGER PRIMARY KEY,
            total REAL DEFAULT 0
        )''')

        if conn.execute('SELECT COUNT(*) FROM balance').fetchone()[0] == 0:
            conn.execute('INSERT INTO balance (id, total) VALUES (1, 0)')
        
        conn.execute('''
        CREATE TABLE IF NOT EXISTS transactions (
            id INTEGER PRIMARY KEY,
            description TEXT,
            amount REAL,
            date TEXT,
            type_id INTEGER,
            FOREIGN KEY(type_id) REFERENCES expense_types(id)
        )''')

        conn.execute('''
        CREATE TABLE IF NOT EXISTS credit_transactions (
            id INTEGER PRIMARY KEY,
            description TEXT,
            amount REAL,
            date TEXT,
            card_id INTEGER,
            FOREIGN KEY(card_id) REFERENCES expense_types(id)
        )''')
        # Kredi kartları tablosu
        conn.execute('''
        CREATE TABLE IF NOT EXISTS credit_cards (
            id INTEGER PRIMARY KEY,
            card_name TEXT,
            credit_limit REAL,
            current_balance REAL DEFAULT 0
        )''')

        

        # Borçlar tablosuna "ödenen taksit sayısı" sütununu ekliyoruz
        conn.execute('''
        CREATE TABLE IF NOT EXISTS debts (
            id INTEGER PRIMARY KEY,
            description TEXT,
            total_amount REAL,
            remaining_amount REAL,
            installment_count INTEGER,
            installment_amount REAL,
            paid_installments INTEGER DEFAULT 0,  -- Ödenen taksit sayısı
            due_date TEXT,
            card_id INTEGER,
            FOREIGN KEY(card_id) REFERENCES credit_cards(id)
        )''')
        # Ödemeler tablosu (kredi kartı veya borç ödemeleri)
        conn.execute('''
        CREATE TABLE IF NOT EXISTS payments (
            id INTEGER PRIMARY KEY,
            amount REAL,
            payment_date TEXT,
            credit_card_id INTEGER,
            debt_id INTEGER,
            FOREIGN KEY(credit_card_id) REFERENCES credit_cards(id),
            FOREIGN KEY(debt_id) REFERENCES debts(id)
        )''')




        conn.commit()


def safe_float(value):
    try:
        return float(value)
    except (ValueError, TypeError):
        return 0.0  # Eğer dönüşüm başarısız olursa 0.0 döndür
# Ana sayfa








@app.route('/')
def index():
    conn = connect_db()
    types = conn.execute('SELECT * FROM expense_types').fetchall()
    
    # Fetch the latest transactions
    transactions = conn.execute('''
        SELECT t.*, e.type_name 
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        ORDER BY t.date DESC LIMIT 10
    ''').fetchall()
    
    balance = conn.execute('SELECT total FROM balance WHERE id = 1').fetchone()['total']
    conn.close()

    # Format transactions for display
    formatted_transactions = [{
        'id': txn['id'],
        'description': txn['description'],
        'amount': "{:+,.2f}".format(float(txn['amount'])),
        'date': datetime.strptime(txn['date'], '%Y-%m-%d').strftime('%d %B %Y'),
        'type_name': txn['type_name'],
        'class': 'positive' if txn['amount'] >= 0 else 'negative'
    } for txn in transactions]

    # Render index with updated transactions
    return render_template('index.html', balance="{:,.2f}".format(balance), transactions=formatted_transactions, types=types)

@app.route('/pay_installment/<int:debt_id>', methods=['POST'])
def pay_installment(debt_id):
    conn = connect_db()

    # Borç bilgilerini alıyoruz
    debt = conn.execute('SELECT remaining_amount, installment_amount, paid_installments, installment_count, card_id FROM debts WHERE id = ?', (debt_id,)).fetchone()

    if debt:
        remaining_amount = debt['remaining_amount']
        installment_amount = debt['installment_amount']
        paid_installments = debt['paid_installments']
        installment_count = debt['installment_count']
        card_id = debt['card_id']

        # Eğer tüm taksitler ödenmemişse ödeme işlemi yapıyoruz
        if paid_installments < installment_count:
            new_remaining_amount = remaining_amount - installment_amount
            new_paid_installments = paid_installments + 1

            # Borcun kalan miktarını ve ödenen taksit sayısını güncelle
            conn.execute('UPDATE debts SET remaining_amount = ?, paid_installments = ? WHERE id = ?', 
                         (new_remaining_amount, new_paid_installments, debt_id))

            # Kredi kartının mevcut bakiyesini güncelle
            conn.execute('UPDATE credit_cards SET current_balance = current_balance - ? WHERE id = ?', 
                         (installment_amount, card_id))

            conn.commit()
            flash('Taksit başarıyla ödendi.')
        else:
            flash('Tüm taksitler ödenmiş.')

    conn.close()
    return redirect('/credit_debt')





@app.route('/add_credit_card', methods=['POST'])
def add_credit_card():
    card_name = request.form['card_name']
    credit_limit = float(request.form['credit_limit'])  # limit yerine credit_limit

    conn = connect_db()
    conn.execute('INSERT INTO credit_cards (card_name, credit_limit, current_balance) VALUES (?, ?, 0)', (card_name, credit_limit))
    conn.commit()
    conn.close()

    flash('Kredi kartı başarıyla eklendi.')
    return redirect('/')


@app.route('/add_debt', methods=['POST'])
def add_debt():
    description = request.form['description']
    total_amount = float(request.form['total_amount'])
    installment_count = int(request.form['installment_count'])
    due_date = request.form['due_date']
    card_id = int(request.form['card_id'])

    # Taksit başına miktarı hesaplıyoruz
    installment_amount = total_amount / installment_count

    conn = connect_db()

    # Kredi kartı bilgilerini alıyoruz
    card = conn.execute('SELECT credit_limit, current_balance FROM credit_cards WHERE id = ?', (card_id,)).fetchone()
    if not card:
        flash('Seçilen kredi kartı bulunamadı.')
        return redirect('/credit_debt')

    # Kartın mevcut bakiyesine göre borç eklemeye çalışıyoruz
    current_balance = card['current_balance']
    credit_limit = card['credit_limit']

    if current_balance + total_amount > credit_limit:
        flash('Borç kart limitini aşıyor! Lütfen borç miktarını düşürün.')
        return redirect('/credit_debt')

    # Borcu ekliyoruz ve kartın mevcut bakiyesini güncelliyoruz
    conn.execute('INSERT INTO debts (description, total_amount, remaining_amount, installment_count, installment_amount, due_date, card_id) VALUES (?, ?, ?, ?, ?, ?, ?)',
                 (description, total_amount, total_amount, installment_count, installment_amount, due_date, card_id))

    # Kartın mevcut bakiyesini güncelle
    conn.execute('UPDATE credit_cards SET current_balance = current_balance + ? WHERE id = ?', (total_amount, card_id))

    conn.commit()
    conn.close()

    flash('Borç başarıyla eklendi.')
    return redirect('/credit_debt')



from datetime import datetime
import locale

from datetime import datetime



# Taksit tarihlerini hesaplayan fonksiyon
def calculate_due_dates(start_date, installment_count):
    due_dates = []
    
    try:
        # Tarihi '%Y-%m-%d' formatında çözümle
        start_date_obj = datetime.strptime(start_date, '%Y-%m-%d')
    except ValueError:
        flash('Geçersiz bir tarih formatı tespit edildi. Varsayılan olarak bugünün tarihi kullanılacak.', 'error')
        start_date_obj = datetime.now()

    for i in range(installment_count):
        next_due_date = start_date_obj + timedelta(days=(i+1) * 30)  # Her taksit için 30 gün ekleniyor
        # Sonucu '%d %B %Y' formatında kaydediyoruz (örn: '14 Aralık 2024')
        due_dates.append(next_due_date.strftime('%d %B %Y'))

    return due_dates


def format_turk(date_string):
    # Tarih formatını "%Y-%m-%d" şeklinde çözümleyip, "%d %B %Y" formatına dönüştürür.
    turkish_months = {
        "January": "Ocak",
        "February": "Şubat",
        "March": "Mart",
        "April": "Nisan",
        "May": "Mayıs",
        "June": "Haziran",
        "July": "Temmuz",
        "August": "Ağustos",
        "September": "Eylül",
        "October": "Ekim",
        "November": "Kasım",
        "December": "Aralık"
    }

    try:
        date_obj = datetime.strptime(date_string, '%Y-%m-%d')
        formatted_date = date_obj.strftime('%d %B %Y')

        for eng_month, tr_month in turkish_months.items():
            formatted_date = formatted_date.replace(eng_month, tr_month)

        return formatted_date
    except ValueError:
        return "Geçersiz tarih"




@app.route('/pay_credit_card', methods=['POST'])
def pay_credit_card():
    card_id = int(request.form['card_id'])
    amount = float(request.form['amount'])
    payment_date = request.form['payment_date']

    conn = connect_db()
    conn.execute('INSERT INTO payments (amount, payment_date, credit_card_id) VALUES (?, ?, ?)', (amount, payment_date, card_id))
    conn.execute('UPDATE credit_cards SET current_balance = current_balance - ? WHERE id = ?', (amount, card_id))
    conn.commit()
    conn.close()

    flash('Kredi kartı borcu başarıyla ödendi.')
    return redirect('/')

# Ödenen taksit sayısını artırmak ve kalan taksit sayısını azaltmak için mevcut fonksiyon
@app.route('/toggle_installment/<int:debt_id>/<int:installment_index>', methods=['POST'])
def toggle_installment(debt_id, installment_index):
    conn = connect_db()

    # Borcun bilgilerini alıyoruz
    debt = conn.execute('SELECT remaining_amount, installment_amount, paid_installments, installment_count, card_id FROM debts WHERE id = ?', (debt_id,)).fetchone()

    if debt:
        remaining_amount = debt['remaining_amount']
        installment_amount = debt['installment_amount']
        paid_installments = debt['paid_installments']
        installment_count = debt['installment_count']
        card_id = debt['card_id']

        # Taksit ödendiyse geri al (tik kaldır), ödenmediyse tik at (taksit öde)
        if installment_index < paid_installments:
            # Taksiti geri alıyoruz (tik kaldırılıyor)
            new_remaining_amount = remaining_amount + installment_amount
            new_paid_installments = paid_installments - 1
        else:
            # Taksiti ödüyoruz (tik atılıyor)
            new_remaining_amount = remaining_amount - installment_amount
            new_paid_installments = paid_installments + 1

        # Borcun kalan miktarını ve ödenen taksit sayısını güncelle
        conn.execute('UPDATE debts SET remaining_amount = ?, paid_installments = ? WHERE id = ?', 
                     (new_remaining_amount, new_paid_installments, debt_id))

        # Kalan taksit sayısını güncelle (total - paid_installments)
        remaining_installments = installment_count - new_paid_installments

        # Kredi kartının mevcut bakiyesini güncelle
        if installment_index < paid_installments:
            # Geri alındığında kart bakiyesini artır
            conn.execute('UPDATE credit_cards SET current_balance = current_balance + ? WHERE id = ?', 
                         (installment_amount, card_id))
        else:
            # Taksit ödendiğinde kart bakiyesini düş
            conn.execute('UPDATE credit_cards SET current_balance = current_balance - ? WHERE id = ?', 
                         (installment_amount, card_id))

        conn.commit()

    conn.close()
    flash('Taksit durumu güncellendi.')
    return redirect('/credit_debt')



@app.route('/pay_debt', methods=['POST'])
def pay_debt():
    debt_id = request.form['debt_id']
    amount = float(request.form['amount'])
    payment_date = request.form['payment_date']

    conn = connect_db()

    # Mevcut borç bilgilerini al
    debt = conn.execute('SELECT remaining_amount, installment_amount FROM debts WHERE id = ?', (debt_id,)).fetchone()

    if debt:
        remaining_amount = debt['remaining_amount']
        installment_amount = debt['installment_amount']

        # Ödeme yapıldığında kalan borçtan düş
        new_remaining_amount = remaining_amount - installment_amount

        conn.execute('UPDATE debts SET remaining_amount = ? WHERE id = ?', (new_remaining_amount, debt_id))
        conn.commit()

    conn.close()
    flash('Borç ödemesi başarıyla yapıldı.')
    return redirect('/credit_debt')



@app.route('/credit_cards')
def credit_cards():
    conn = connect_db()
    credit_cards = conn.execute('SELECT * FROM credit_cards').fetchall()
    conn.close()
    return render_template('credit_cards.html', credit_cards=credit_cards)

@app.route('/debts')
def debts():
    conn = connect_db()
    debts = conn.execute('SELECT * FROM debts').fetchall()
    conn.close()
    return render_template('debts.html', debts=debts)

@app.route('/credit_debt')
def credit_debt():
    conn = connect_db()

    # Kredi kartlarını ve mevcut bakiye/limit bilgilerini çekiyoruz
    credit_cards = conn.execute('SELECT * FROM credit_cards').fetchall()

    # Borçları ve ilgili kart isimlerini çekiyoruz
    debts = conn.execute('''
        SELECT d.*, c.card_name, c.credit_limit, c.current_balance
        FROM debts d
        JOIN credit_cards c ON d.card_id = c.id
    ''').fetchall()

    conn.close()

    # Kredi kartlarını formatlıyoruz ve float dönüştürme yapıyoruz
    formatted_credit_cards = [{
        'id': card['id'],
        'card_name': card['card_name'],
        'credit_limit': "{:,.2f}".format(safe_float(card['credit_limit'])),
        'current_balance': "{:,.2f}".format(safe_float(card['current_balance'])),
        'remaining_limit': "{:,.2f}".format(safe_float(card['credit_limit']) - safe_float(card['current_balance']))
    } for card in credit_cards]

    # Borçları formatlıyoruz
    formatted_debts = [{
        'id': debt['id'],
        'description': debt['description'],
        'total_amount': "{:,.2f}".format(safe_float(debt['total_amount'])),
        'remaining_amount': "{:,.2f}".format(safe_float(debt['remaining_amount'])),
        'installment_count': debt['installment_count'],
        'installment_amount': "{:,.2f}".format(safe_float(debt['installment_amount'])),
        'paid_installments': debt['paid_installments'],
        # Son ödeme tarihini gün, ay ismi ve yıl formatında gösteriyoruz
        'due_date': format_turk(debt['due_date']),
        'card_name': debt['card_name']
    } for debt in debts]

    return render_template('credit_debt.html', credit_cards=formatted_credit_cards, debts=formatted_debts)

@app.route('/update_credit_limit', methods=['POST'])
def update_credit_limit():
    card_id = request.form['card_id']
    new_limit = float(request.form['new_limit'])

    conn = connect_db()

    # Kart limitini güncelleme
    conn.execute('UPDATE credit_cards SET credit_limit = ? WHERE id = ?', (new_limit, card_id))
    conn.commit()
    conn.close()

    flash('Kredi kartı limiti başarıyla güncellendi.', 'success')
    return redirect('/credit_debt')

@app.route('/delete_credit_card/<int:card_id>', methods=['POST'])
def delete_credit_card(card_id):
    conn = connect_db()

    try:
        # Kredi kartı ile ilişkili olan tüm harcamaları siliyoruz
        conn.execute('DELETE FROM credit_transactions WHERE card_id = ?', (card_id,))
        # Kredi kartı ile ilişkili tüm borçları siliyoruz
        conn.execute('DELETE FROM debts WHERE card_id = ?', (card_id,))
        # Kredi kartının kendisini siliyoruz
        conn.execute('DELETE FROM credit_cards WHERE id = ?', (card_id,))

        conn.commit()
        flash('Kredi kartı ve ilişkili tüm harcamalar başarıyla silindi.', 'success')
    except sqlite3.Error as e:
        conn.rollback()
        flash(f'Bir hata oluştu: {e}', 'error')

    conn.close()
    return redirect('/credit_debt')
@app.route('/delete_debt/<int:debt_id>', methods=['POST'])
def delete_debt(debt_id):
    conn = connect_db()

    try:
        # Borcun bilgilerini çekiyoruz
        debt = conn.execute('SELECT total_amount, remaining_amount, paid_installments, card_id FROM debts WHERE id = ?', (debt_id,)).fetchone()

        if debt:
            total_amount = debt['total_amount']
            remaining_amount = debt['remaining_amount']  # Kalan borç miktarı
            paid_installments = debt['paid_installments']  # Ödenen taksit sayısı
            card_id = debt['card_id']

            # Borcu siliyoruz
            conn.execute('DELETE FROM debts WHERE id = ?', (debt_id,))

            # Ödenen taksitlerin toplamını hesapla
            paid_amount = total_amount - remaining_amount

            # Kalan borcu güncellemek için mevcut bakiyeyi güncelle
            conn.execute('UPDATE credit_cards SET current_balance = current_balance - ? WHERE id = ?', (remaining_amount, card_id))

            conn.commit()
            flash('Borç ve kartın mevcut bakiyesi başarıyla güncellendi.', 'success')
        else:
            flash('Borç bulunamadı.', 'error')
    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Bir hata oluştu: {e}", 'error')
    finally:
        conn.close()

    return redirect('/credit_debt')





# Taksit tarihlerini hesaplayan fonksiyon
def calculate_due_dates(start_date, installment_count):
    due_dates = []
    start_date_obj = datetime.strptime(start_date, '%Y-%m-%d')
    
    for i in range(installment_count):
        next_due_date = start_date_obj + timedelta(days=(i+1) * 30)  # Her taksit için 30 gün ekleniyor
        due_dates.append(next_due_date.strftime('%Y-%m-%d'))  # Yıl-Ay-Gün formatında ekle
    return due_dates



# Harcama raporu sayfası
@app.route('/report')
def report():
    conn = connect_db()
    report_data = conn.execute('''
        SELECT e.type_name, SUM(t.amount) as total_amount
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        GROUP BY e.type_name
        HAVING total_amount < 0
    ''').fetchall()
    conn.close()

    formatted_report = [{
        'type_name': row['type_name'],
        'total_amount': abs(float(row['total_amount']))
    } for row in report_data]

    return render_template('report.html', report_data=formatted_report)

# Anapara ekleme
@app.route('/add_funds', methods=['POST'])
def add_funds():
    description = request.form['description']
    amount = float(request.form['amount'])
    date = request.form['date']
    type_id = request.form['type_id']

    conn = connect_db()
    conn.execute('INSERT INTO transactions (description, amount, date, type_id) VALUES (?, ?, ?, ?)', (description, amount, date, type_id))
    conn.execute('UPDATE balance SET total = total + ? WHERE id = 1', (amount,))
    conn.commit()
    conn.close()

    return redirect('/')

# Harcama ekleme
@app.route('/add_expense', methods=['POST'])
def add_expense():
    description = request.form['description']
    amount = -abs(float(request.form['amount']))  # Harcama negatif olarak kaydedilir
    date = request.form['date']
    type_id = request.form['type_id']

    conn = connect_db()
    conn.execute('INSERT INTO transactions (description, amount, date, type_id) VALUES (?, ?, ?, ?)', (description, amount, date, type_id))
    conn.execute('UPDATE balance SET total = total + ? WHERE id = 1', (amount,))
    conn.commit()
    conn.close()

    return redirect('/')



@app.route('/credit_add_expense', methods=['POST'])
def credit_add_expense():
    description = request.form['description']
    amount = -abs(float(request.form['amount']))  # Harcama negatif olarak kaydedilir
    installment_count = int(request.form['installment_count'])  # Taksit sayısı
    expense_date = request.form['expense_date']  # Harcama tarihi (kullanıcıdan alınır)
    card_id = int(request.form['card_id'])  # Kullanıcının seçtiği kart

    conn = connect_db()

    # Seçilen kartın mevcut bakiyesini ve limitini kontrol et
    card = conn.execute('SELECT credit_limit, current_balance FROM credit_cards WHERE id = ?', (card_id,)).fetchone()
    if not card:
        flash('Seçilen kredi kartı bulunamadı.', 'error')
        return redirect('/credit_debt')

    current_balance = card['current_balance']
    credit_limit = card['credit_limit']

    # Kartın mevcut bakiyesine harcama eklenir
    new_balance = current_balance + abs(amount)

    # Eğer harcama kart limitini aşıyorsa, uyarı ver
    if new_balance > credit_limit:
        flash('Harcama kart limitini aşıyor! Lütfen harcama miktarını düşürün.', 'error')
        return redirect('/credit_debt')

    # Harcama kaydedilir ve kartın mevcut bakiyesi güncellenir
    conn.execute('INSERT INTO credit_transactions (description, amount, date, card_id) VALUES (?, ?, ?,? )', (description, amount, expense_date,card_id))

    # Kartın mevcut bakiyesini güncelle
    conn.execute('UPDATE credit_cards SET current_balance = ? WHERE id = ?', (new_balance, card_id))

    # Borcu ekle (Borçlar tablosuna yeni harcamayı ekle)
    installment_amount = abs(amount) / installment_count
    due_dates = calculate_due_dates(expense_date, installment_count)  # Taksit tarihlerini hesapla

    # Son taksitin ödeme tarihini istediğiniz formatta (gün ay ismi yıl) ayarlıyoruz
    formatted_due_date = due_dates[-1]  # Zaten doğru formatta hesaplanmış durumda

    conn.execute('INSERT INTO debts (description, total_amount, remaining_amount, installment_count, installment_amount, paid_installments, due_date, card_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                 (description, abs(amount), abs(amount), installment_count, installment_amount, 0, formatted_due_date, card_id))

    conn.commit()
    conn.close()

    # Harcama başarıyla eklendiğinde flash mesajı gönder
    flash('Harcama başarıyla eklendi!', 'success')
    return redirect('/credit_debt')










# CSV Dışa Aktarma İşlevi
@app.route('/export_csv')
def export_csv():
    conn = connect_db()
    transactions = conn.execute('''
        SELECT t.*, e.type_name
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        ORDER BY t.date DESC
    ''').fetchall()
    conn.close()

    si = BytesIO()
    cw = csv.writer(si)

    # UTF-8 ile BOM eklemek
    si.write('\ufeff')

    # CSV Başlıkları
    cw.writerow(['ID', 'Açıklama', 'Miktar (TL)', 'Tarih', 'Tür'])

    for txn in transactions:
        formatted_amount = "{:+,.2f}".format(float(txn['amount']))
        formatted_date = datetime.strptime(txn['date'], '%Y-%m-%d').strftime('%d-%m-%Y')
        cw.writerow([txn['id'], txn['description'], formatted_amount, formatted_date, txn['type_name']])

    output = make_response(si.getvalue())
    output.headers["Content-Disposition"] = "attachment; filename=transactions.csv"
    output.headers["Content-type"] = "text/csv; charset=utf-8"
    return output

# Harcama raporunu Word formatında dışa aktarma
@app.route('/export_word')
def export_word():
    conn = connect_db()
    report_data = conn.execute('''
        SELECT e.type_name, SUM(t.amount) as total_amount
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        GROUP BY e.type_name
        HAVING total_amount < 0
    ''').fetchall()
    conn.close()

    # Word dokümanı oluşturma
    doc = Document()
    doc.add_heading('Harcama Raporu', 0)

    # Tablo ekliyoruz
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Harcama Türü'
    hdr_cells[1].text = 'Toplam Harcama (TL)'

    for row in report_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row['type_name']
        row_cells[1].text = "{:,.2f}".format(abs(float(row['total_amount']))) + ' TL'

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    response = make_response(f.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=harcama_raporu.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return response

# İşlem silme
@app.route('/delete_transaction/<int:transaction_id>', methods=['POST'])
def delete_transaction(transaction_id):
    conn = connect_db()
    
    try:
        # Silinecek işlem bilgisini alıyoruz
        transaction = conn.execute('SELECT amount FROM transactions WHERE id = ?', (transaction_id,)).fetchone()
        
        if transaction:
            amount = transaction['amount']
            # İşlemi veritabanından siliyoruz
            conn.execute('DELETE FROM transactions WHERE id = ?', (transaction_id,))
            # Güncel bakiyeyi güncelliyoruz
            conn.execute('UPDATE balance SET total = total - ? WHERE id = 1', (amount,))
            conn.commit()
            flash('İşlem başarıyla silindi.', 'success')
        else:
            flash('İşlem bulunamadı.', 'error')
            
    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Bir hata oluştu: {e}", 'error')
    finally:
        conn.close()
    
    # `index` route'una yönlendirme yapıyoruz
    return redirect('/')
@app.route('/all_transactions')
def all_transactions():
    conn = connect_db()
    transactions = conn.execute('''
        SELECT t.*, e.type_name 
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        ORDER BY t.date DESC
    ''').fetchall()
    conn.close()

    # Format işlemlerini şablon için hazırlayın
    formatted_transactions = [{
        'id': txn['id'],
        'description': txn['description'],
        'amount': "{:+,.2f}".format(float(txn['amount'])),
        'date': datetime.strptime(txn['date'], '%Y-%m-%d').strftime('%d %B %Y'),
        'type_name': txn['type_name'],
        'class': 'positive' if txn['amount'] >= 0 else 'negative'
    } for txn in transactions]

    return render_template('all_transactions.html', transactions=formatted_transactions)
@app.route('/load_transactions')
def load_transactions():
    page = int(request.args.get('page', 1))
    limit = int(request.args.get('limit', 10))
    offset = (page - 1) * limit
    
    conn = connect_db()
    transactions = conn.execute('''
        SELECT t.id, t.description, t.amount, t.date, e.type_name
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        ORDER BY t.date DESC
        LIMIT ? OFFSET ?
    ''', (limit, offset)).fetchall()
    conn.close()

    # Format the data to be JSON serializable
    transaction_data = [{
        'id': txn['id'],
        'description': txn['description'],
        'amount': "{:+,.2f}".format(float(txn['amount'])),
        'date': datetime.strptime(txn['date'], '%Y-%m-%d').strftime('%d %B %Y'),
        'type_name': txn['type_name']
    } for txn in transactions]

    return jsonify(transaction_data)

# Filtreleme fonksiyonu
@app.route('/filter', methods=['GET', 'POST'])
def filter_transactions():
    conn = connect_db()

    query = '''
    SELECT t.*, e.type_name
    FROM transactions t
    LEFT JOIN expense_types e ON t.type_id = e.id
    WHERE 1=1
'''
    transactions = conn.execute(query).fetchall()
    filters = []

    # Eğer POST isteği geldiyse filtreleme yapılacak
    if request.method == 'POST':
        start_date = request.form.get('start_date')  # Tarih YYYY-MM-DD formatında
        end_date = request.form.get('end_date')  # Tarih YYYY-MM-DD formatında
        type_id = request.form.get('type_id')

        # Tarih aralığı kontrolü
        if start_date:
            query += ' AND t.date >= ?'
            filters.append(start_date)

        if end_date:
            query += ' AND t.date <= ?'
            filters.append(end_date)

        type_id = request.form.get('type_id')

# 'all' değilse type_id'yi integer'a dönüştür
        if type_id and type_id != 'all':
            type_id = int(type_id)  # string olan type_id'yi int'e çeviriyoruz
            query += ' AND t.type_id = ?'
            filters.append(type_id)


    # Sorgu sonucunu çalıştırmadan önce sorguyu çıktı olarak yazdırıyoruz
    print(f"Final SQL Query: {query}")
    print(f"Filters: {filters}")

    query += ' ORDER BY t.date DESC'
    transactions = conn.execute(query, filters).fetchall()

    # Harcama türleri
    types = conn.execute('SELECT * FROM expense_types').fetchall()

    # Güncel bakiye
    balance = conn.execute('SELECT total FROM balance WHERE id = 1').fetchone()['total']
    conn.close()
    

    # Eğer sonuç yoksa kontrol edin
    if not transactions:
        flash("Seçilen filtrelere göre sonuç bulunamadı", "info")

    formatted_transactions = [{
        'id': txn['id'],
        'description': txn['description'],
        'amount': "{:+,.2f}".format(float(txn['amount'])),
        'date': datetime.strptime(txn['date'], '%Y-%m-%d').strftime('%d %B %Y'),
        'type_name': txn['type_name'],
        'class': 'positive' if txn['amount'] >= 0 else 'negative'
        
    } for txn in transactions]
    print(f"Transactions: {formatted_transactions}")
    

    return render_template('filter.html', transactions=formatted_transactions, types=types, balance=balance)


# Günlük harcama raporu sayfası
@app.route('/daily_report', methods=['GET'])
def daily_report():
    conn = connect_db()

    # Tarih seçenekleri
    today = datetime.now().date()
    yesterday = today - timedelta(days=1)

    # Dropdown'dan gelen değeri alıyoruz
    selected_date = request.args.get('selected_date', 'today')
    specific_date = request.args.get('specific_date')  # Spesifik tarih alanı

    if selected_date == 'today':
        selected_day = today
        date_text = 'Bugün'
    elif selected_date == 'yesterday':
        selected_day = yesterday
        date_text = 'Dün'
    elif selected_date == 'specific' and specific_date:
        try:
            selected_day = datetime.strptime(specific_date, '%Y-%m-%d').date()
            date_text = selected_day.strftime('%d %B %Y')
        except ValueError:
            flash('Geçersiz tarih seçimi!', 'error')
            return redirect('/daily_report')
    else:
        flash('Geçersiz seçim!', 'error')
        return redirect('/daily_report')

    # Veritabanından harcamaları alıyoruz
    query = '''
    SELECT description, amount, date FROM transactions
    WHERE date(date) = ?
    '''
    transactions = conn.execute(query, (selected_day,)).fetchall()
    conn.close()

    # Şablona verileri gönderiyoruz
    return render_template('daily_report.html', data=transactions, selected_date=selected_date, specific_date=specific_date, date_text=date_text)


@app.route('/statistics', methods=['GET', 'POST'])
def statistics():
    conn = connect_db()

    # Formdan gelen yıl ve ay bilgileri (örn: 2024 ve 01)
    selected_year = request.form.get('year') if request.method == 'POST' else None
    selected_month = request.form.get('month') if request.method == 'POST' else None

    # Eğer yıl ve ay seçilmediyse, varsayılan olarak mevcut yıl ve ayı seç
    if not selected_year:
        selected_year = datetime.now().strftime('%Y')  # Varsayılan yıl
    if not selected_month:
        selected_month = datetime.now().strftime('%m')  # Varsayılan ay

    # Seçilen yıl ve aya göre harcama türü istatistiklerini getiriyoruz
    statistics_data = conn.execute('''
        SELECT e.type_name, SUM(t.amount) AS total_spent
        FROM transactions t
        LEFT JOIN expense_types e ON t.type_id = e.id
        WHERE strftime('%Y', t.date) = ? AND strftime('%m', t.date) = ?
        GROUP BY e.type_name
    ''', (selected_year, selected_month)).fetchall()

    # Toplam harcamayı hesaplama (Tüm harcamaların toplamı)
    total_expense = conn.execute('''
        SELECT SUM(amount) FROM transactions WHERE amount < 0
    ''').fetchone()[0] or 0  # Harcama negatifse, toplam negatif olmasın diye sıfır döndür

    # Yıllar ve aylar için benzersiz değerleri alıyoruz (Dropdown menüler için)
    available_years = conn.execute('''
        SELECT DISTINCT strftime('%Y', t.date) AS year
        FROM transactions t
        ORDER BY year DESC
    ''').fetchall()

    available_months = conn.execute('''
        SELECT DISTINCT strftime('%m', t.date) AS month
        FROM transactions t
        ORDER BY month ASC
    ''').fetchall()

    conn.close()

    # Şablona gönderilecek veriler
    formatted_statistics = [{
        'type_name': row['type_name'],
        'total_spent': abs(row['total_spent'])  # Negatif değerleri pozitif yapıyoruz
    } for row in statistics_data]

    # Ay isimlerini Türkçe olarak formatlama
    available_months = [{
        'month': m['month'],
        'formatted': datetime.strptime(m['month'], '%m').strftime('%B')  # Ay ismini Türkçe olarak al
    } for m in available_months]

    # Şablona toplam harcama ekleniyor
    return render_template('statistics.html', 
                           statistics=formatted_statistics, 
                           years=available_years, 
                           months=available_months, 
                           selected_year=selected_year, 
                           selected_month=selected_month,
                           total_expense=total_expense)


if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='0.0.0.0', port=8000)
